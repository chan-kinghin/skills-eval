"""Document preprocessing for SkillEval.

Converts input files (PDF, DOCX, XLSX, and plain text) into text strings
that can be inlined into LLM prompts. Each file type has a dedicated
extractor; unsupported types fall back to raw text reading.
"""

from __future__ import annotations

from pathlib import Path

# Supported file extensions and their categories
TEXT_EXTENSIONS = {".txt", ".md", ".json", ".csv", ".xml", ".html", ".yaml", ".yml", ".tsv"}
PDF_EXTENSIONS = {".pdf"}
DOCX_EXTENSIONS = {".docx"}
XLSX_EXTENSIONS = {".xlsx", ".xls"}


def extract_text(file_path: Path) -> str:
    """Extract text content from a file, dispatching by extension.

    Supports: plain text, PDF (pdfplumber), DOCX (python-docx), XLSX (openpyxl).
    Raises RuntimeError with install instructions if an optional dependency is missing.
    """
    suffix = file_path.suffix.lower()

    if suffix in PDF_EXTENSIONS:
        return _extract_pdf(file_path)
    elif suffix in DOCX_EXTENSIONS:
        return _extract_docx(file_path)
    elif suffix in XLSX_EXTENSIONS:
        return _extract_xlsx(file_path)
    elif suffix in TEXT_EXTENSIONS or _is_likely_text(file_path):
        return file_path.read_text(encoding="utf-8")
    else:
        # Last resort: try reading as text, fail gracefully
        try:
            return file_path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            raise ValueError(
                f"Cannot read '{file_path.name}': unsupported binary format. "
                f"Supported formats: {', '.join(sorted(TEXT_EXTENSIONS | PDF_EXTENSIONS | DOCX_EXTENSIONS | XLSX_EXTENSIONS))}"
            )


def _extract_pdf(file_path: Path) -> str:
    """Extract text from a PDF using pdfplumber."""
    try:
        import pdfplumber
    except ImportError:
        raise RuntimeError(
            "PDF support requires pdfplumber. Install it with: pip install pdfplumber"
        )

    pages = []
    with pdfplumber.open(file_path) as pdf:
        for i, page in enumerate(pdf.pages, 1):
            text = page.extract_text()
            if text:
                pages.append(f"[Page {i}]\n{text}")

            # Also extract tables if present
            tables = page.extract_tables()
            for j, table in enumerate(tables, 1):
                if table:
                    table_text = _format_table(table)
                    pages.append(f"[Page {i} - Table {j}]\n{table_text}")

    if not pages:
        raise ValueError(
            f"Could not extract text from '{file_path.name}'. The PDF may be scanned/image-based."
        )

    return "\n\n".join(pages)


def _extract_docx(file_path: Path) -> str:
    """Extract text from a Word document using python-docx."""
    try:
        import docx
    except ImportError:
        raise RuntimeError(
            "DOCX support requires python-docx. Install it with: pip install python-docx"
        )

    doc = docx.Document(str(file_path))
    parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Extract tables
    for i, table in enumerate(doc.tables, 1):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        if rows:
            table_text = _format_table(rows)
            parts.append(f"[Table {i}]\n{table_text}")

    if not parts:
        raise ValueError(
            f"Could not extract text from '{file_path.name}'. The document appears to be empty."
        )

    return "\n\n".join(parts)


def _extract_xlsx(file_path: Path) -> str:
    """Extract text from an Excel file using openpyxl."""
    try:
        import openpyxl
    except ImportError:
        raise RuntimeError("XLSX support requires openpyxl. Install it with: pip install openpyxl")

    wb = openpyxl.load_workbook(str(file_path), read_only=True, data_only=True)
    parts = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(cell) if cell is not None else "" for cell in row]
            if any(c.strip() for c in cells):  # skip empty rows
                rows.append(cells)

        if rows:
            table_text = _format_table(rows)
            parts.append(f"[Sheet: {sheet_name}]\n{table_text}")

    wb.close()

    if not parts:
        raise ValueError(
            f"Could not extract data from '{file_path.name}'. The workbook appears to be empty."
        )

    return "\n\n".join(parts)


def _format_table(rows: list[list[str | None]]) -> str:
    """Format a table as a pipe-delimited text representation.

    This is a simple, LLM-friendly format:
    | Header1 | Header2 | Header3 |
    | ------- | ------- | ------- |
    | Value1  | Value2  | Value3  |
    """
    if not rows:
        return ""

    # Normalize cells
    normalized = []
    for row in rows:
        normalized.append([str(cell).strip() if cell else "" for cell in row])

    # Build pipe-delimited table
    lines = []
    for i, row in enumerate(normalized):
        line = "| " + " | ".join(row) + " |"
        lines.append(line)
        # Add separator after first row (header)
        if i == 0:
            sep = "| " + " | ".join("---" for _ in row) + " |"
            lines.append(sep)

    return "\n".join(lines)


def _is_likely_text(file_path: Path) -> bool:
    """Heuristic check: read first 512 bytes to see if the file looks like text."""
    try:
        with open(file_path, "rb") as f:
            chunk = f.read(512)
        # If it decodes as UTF-8 without errors, it's probably text
        chunk.decode("utf-8")
        return True
    except (UnicodeDecodeError, OSError):
        return False


def format_input_files(input_files: list[Path]) -> str:
    """Convert a list of input files into an LLM-ready text prompt.

    Each file is extracted to text and wrapped with delimiters:
    --- File: invoice.pdf ---
    [Page 1]
    ACME Corp Invoice #1234
    ...
    --- End File ---
    """
    parts = []
    for f in input_files:
        content = extract_text(f)
        parts.append(f"--- File: {f.name} ---\n{content}\n--- End File ---")
    return "\n\n".join(parts)


async def extract_text_async(file_path: Path) -> str:
    """Async wrapper for extract_text that runs I/O in a thread."""
    import asyncio

    return await asyncio.to_thread(extract_text, file_path)


async def format_input_files_async(input_files: list[Path]) -> str:
    """Async wrapper for format_input_files that runs I/O in a thread."""
    import asyncio

    return await asyncio.to_thread(format_input_files, input_files)


def input_descriptions(input_files: list[Path], max_chars: int = 200) -> str:
    """Generate short descriptions of input files for skill generation prompts.

    Extracts text and shows a preview (first max_chars characters).
    """
    parts = []
    for f in input_files:
        try:
            content = extract_text(f)
            preview = content[:max_chars].replace("\n", " ")
            if len(content) > max_chars:
                preview += "..."
            suffix = f.suffix.lower()
            type_label = {
                ".pdf": "PDF",
                ".docx": "Word",
                ".xlsx": "Excel",
                ".json": "JSON",
                ".csv": "CSV",
            }.get(suffix, "text")
            parts.append(f"- {f.name} ({type_label}): {preview}")
        except (ValueError, RuntimeError) as e:
            parts.append(f"- {f.name}: [could not preview: {e}]")
    return "\n".join(parts)
